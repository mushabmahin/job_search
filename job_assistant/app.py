"""
Job Application Assistant — Deep Agent powered by LangChain deepagents + Tavily + OpenAI
Run:  streamlit run app.py
"""

import os
import io
import json
import re
from typing import Literal, Dict, Any, List

import streamlit as st
import pandas as pd

from langchain_core.tools import tool
from langchain.chat_models import init_chat_model
from tavily import TavilyClient
from deepagents import create_deep_agent

import pypdf
import docx
from dotenv import load_dotenv

# Load env variables from .env
load_dotenv()


# ─────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────
st.set_page_config(
    page_title="Job Application Assistant",
    page_icon="💼",
    layout="wide",
)

st.markdown("""
<style>
    .stApp { background-color: #0d1117; color: #c9d1d9; }
    .block-container { padding-top: 2rem; }
    h1 { color: #58a6ff; }
    h2, h3 { color: #79c0ff; }
    .stButton>button {
        background: linear-gradient(135deg, #238636, #2ea043);
        color: white; border: none; border-radius: 6px;
        padding: 0.5rem 2rem; font-weight: 600;
    }
    .stButton>button:hover { background: linear-gradient(135deg, #2ea043, #3fb950); }
    .stDownloadButton>button {
        background: linear-gradient(135deg, #1f6feb, #388bfd);
        color: white; border: none; border-radius: 6px;
    }
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea {
        background-color: #161b22;
        color: #c9d1d9;
        border: 1px solid #30363d;
        border-radius: 6px;
    }
    table { width: 100%; border-collapse: collapse; }
    th { background-color: #161b22; color: #79c0ff; padding: 10px; text-align: left; border-bottom: 2px solid #30363d; }
    td { padding: 10px; border-bottom: 1px solid #21262d; color: #c9d1d9; }
    tr:hover td { background-color: #161b22; }
    a { color: #58a6ff; text-decoration: none; }
    a:hover { text-decoration: underline; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────
for key, default in [
    ("jobs_df", None),
    ("cover_doc", None),
    ("last_error", ""),
    ("raw_final", ""),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ─────────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────────
st.title("💼 Job Application Assistant")
st.markdown("*Powered by LangChain Deep Agents · Tavily Search · OpenAI GPT-4o-mini*")
st.divider()

# ─────────────────────────────────────────────────
# UI — Inputs
# ─────────────────────────────────────────────────
c0, c1, c2 = st.columns([2, 1, 1])
with c0:
    uploaded = st.file_uploader(
        "📄 Upload your resume (PDF / DOCX / TXT)",
        type=["pdf", "docx", "txt"],
    )
with c1:
    target_title = st.text_input("🎯 Target title", "Senior Machine Learning Engineer")
with c2:
    target_location = st.text_input("📍 Target location(s)", "Bangalore OR Remote")

skills_hint = st.text_area(
    "🔧 Add / override skills (optional)",
    "",
    placeholder="Python, PyTorch, LLMs, RAG, Azure, vLLM, FastAPI",
    height=80,
)

# API Keys configuration (loaded from environment / .env file)
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Pre-fill from environment
    env_openai_key = os.environ.get("OPENAI_API_KEY", "").strip()
    env_tavily_key = os.environ.get("TAVILY_API_KEY", "").strip()
    
    openai_key_input = st.text_input(
        "🔑 OpenAI API Key",
        value=env_openai_key,
        type="password",
        placeholder="sk-proj-..."
    )
    
    tavily_key_input = st.text_input(
        "🔑 Tavily API Key",
        value=env_tavily_key,
        type="password",
        placeholder="tvly-dev-..."
    )

    openai_model = st.selectbox(
        "OpenAI Model",
        ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"],
        index=0,
    )
    st.divider()
    st.markdown("**How it works**")
    st.markdown("""
1. Upload your resume (PDF/DOCX/TXT)
2. Set your target title & location
3. Optionally add skills to highlight
4. Click **Run** — the agent will:
   - Search for 5 matching live jobs
   - Generate a tailored cover letter per job
5. Download the cover letters as `.docx`
    """)
    st.divider()
    st.markdown("Get a free Tavily key → [app.tavily.com](https://app.tavily.com)")


# ─────────────────────────────────────────────────
# FILE HELPERS
# ─────────────────────────────────────────────────
def extract_text(file) -> str:
    if not file:
        return ""
    name = file.name.lower()
    if name.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        pdf = pypdf.PdfReader(io.BytesIO(file.read()))
        return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if name.endswith(".docx"):
        d = docx.Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in d.paragraphs)
    return ""


def md_to_docx(md_text: str) -> bytes:
    doc = docx.Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            doc.add_heading(line.lstrip("#").strip(), level=level)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ─────────────────────────────────────────────────
# DATA EXTRACTION
# ─────────────────────────────────────────────────
def normalize_jobs(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normed = []
    for it in items:
        if not isinstance(it, dict):
            continue
        lower_map = {str(k).strip().lower(): it[k] for k in it.keys()}
        company  = str(lower_map.get("company",  "") or "").strip()
        title    = str(lower_map.get("title",    "") or "").strip()
        location = str(lower_map.get("location", "") or "").strip()
        link     = str(lower_map.get("link",     "") or "").strip()
        why_fit  = str(
            lower_map.get("why_fit", lower_map.get("good match", "")) or ""
        ).strip()
        if not link:
            continue
        normed.append({
            "Company":    company  or "—",
            "Title":      title    or "—",
            "Location":   location or "—",
            "Link":       link,
            "Good Match": "✅ Yes" if why_fit else "—",
        })
    return normed[:5]


def extract_jobs_from_text(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    pattern = r"<JOBS>\s*(?:```[\w-]*\s*)?(\[.*?\])\s*(?:```)?\s*</JOBS>"
    m = re.search(pattern, text, flags=re.S | re.I)
    if not m:
        return []
    raw = m.group(1).strip().strip("`").strip()
    try:
        obj = json.loads(raw)
        return obj if isinstance(obj, list) else []
    except Exception:
        try:
            salvaged = re.sub(r"(?<!\\)'", '"', raw)
            obj = json.loads(salvaged)
            return obj if isinstance(obj, list) else []
        except Exception:
            st.session_state.last_error = f"JSON parse failed:\n{raw[:1200]}"
            return []


# ─────────────────────────────────────────────────
# TOOL
# Use a mutable dict so the @tool closure always reads the latest key
# without needing to re-import the module (which causes duplicate widget IDs).
# ─────────────────────────────────────────────────
_runtime = {"tavily_key": ""}


@tool
def internet_search(
    query: str,
    max_results: int = 5,
    topic: Literal["general", "news", "finance"] = "general",
) -> str:
    """Search the web for current information. Returns a JSON string of results."""
    key = _runtime["tavily_key"]
    if not key:
        raise RuntimeError("TAVILY_API_KEY is not set.")
    client = TavilyClient(api_key=key)
    # .search() returns a dict: {"query":..., "results":[...], "answer":...}
    # deepagents expects tool outputs to be strings, so we extract the
    # results list and serialise it to JSON.
    response = client.search(
        query=query,
        max_results=max_results,
        topic=topic,
    )
    results = response.get("results", [])
    # Trim to essential fields to keep context short
    trimmed = [
        {
            "title":   r.get("title", ""),
            "url":     r.get("url", ""),
            "content": r.get("content", "")[:400],
        }
        for r in results
    ]
    return json.dumps(trimmed, ensure_ascii=False)


# ─────────────────────────────────────────────────
# AGENT PROMPTS
# ─────────────────────────────────────────────────
INSTRUCTIONS = (
    "You are a job application assistant. Do exactly two things in order:\n\n"
    "1) Use the internet_search tool to find exactly 5 CURRENT, REAL job postings that "
    "match the user's target title, locations, and skills. "
    "Return them ONLY as a JSON array wrapped in these exact tags:\n"
    "<JOBS>\n"
    '[{"company":"...","title":"...","location":"...","link":"https://...","Good Match":"one sentence why this fits"},'
    " ... five objects total]\n"
    "</JOBS>\n"
    "Rules: valid JSON only (no comments, no trailing commas), real clickable links, no duplicate companies.\n\n"
    "2) For EACH of the 5 jobs, write a personalised cover letter (≤150 words) with a subject line, "
    "appended to cover_letters.md under a markdown heading per job.\n\n"
    "Do NOT invent fake companies or links. "
    "Prefer company career pages, LinkedIn, Lever, Greenhouse, or Indeed."
)

JOB_SEARCH_PROMPT = (
    "You specialise in finding real, current job postings. "
    "Search and select 5 postings matching the user's criteria. "
    "Output ONLY the <JOBS> block — no text before or after it:\n"
    "<JOBS>\n"
    '[{"company":"Acme","title":"ML Engineer","location":"Remote","link":"https://...","Good Match":"Strong Python and LLM experience match."},'
    '{"company":"Beta Corp","title":"AI Engineer","location":"Bangalore","link":"https://...","Good Match":"RAG and PyTorch align well."},'
    '{"company":"Gamma Inc","title":"Senior ML Engineer","location":"Hyderabad","link":"https://...","Good Match":"Azure ML stack experience."},'
    '{"company":"Delta Ltd","title":"NLP Engineer","location":"Remote","link":"https://...","Good Match":"LLM fine-tuning background fits."},'
    '{"company":"Epsilon AI","title":"ML Platform Engineer","location":"Bangalore","link":"https://...","Good Match":"FastAPI and vLLM experience."}]'
    "\n</JOBS>"
)

COVER_LETTER_PROMPT = (
    "You are an expert cover letter writer. "
    "For each job in the found list, write a subject line and a tight cover letter (≤150 words) "
    "that connects the user's specific skills and experience from their resume to this exact role. "
    "Append everything to cover_letters.md under a ## heading per job. "
    "Be specific, not generic. Avoid fluff."
)


# ─────────────────────────────────────────────────
# BUILD AGENT
# ─────────────────────────────────────────────────
def build_agent(openai_key: str, openai_model_name: str):
    os.environ["OPENAI_API_KEY"] = openai_key
    # Use init_chat_model so deepagents gets a BaseChatModel object — this avoids
    # the "expected string or bytes-like object, got list" error that occurs when
    # deepagents tries to concatenate instructions with the base prompt and finds
    # a list instead of a string due to argument mis-ordering in older API versions.
    llm = init_chat_model(f"openai:{openai_model_name}")
    subagents = [
        {
            "name": "job-search-agent",
            "description": "Searches the web and finds 5 real, current job postings matching the criteria.",
            "system_prompt": JOB_SEARCH_PROMPT,
            "tools": [internet_search],
        },
        {
            "name": "cover-letter-writer-agent",
            "description": "Writes tailored cover letters for each found job and saves them to cover_letters.md.",
            "system_prompt": COVER_LETTER_PROMPT,
        },
    ]
    return create_deep_agent(
        model=llm,
        tools=[internet_search],
        system_prompt=INSTRUCTIONS,
        subagents=subagents,
    )


def make_task_prompt(resume_text: str, skills: str, title: str, location: str) -> str:
    skill_line = f"\nPrioritize these skills: {skills.strip()}." if skills.strip() else ""
    return (
        f"Target job title: {title}\n"
        f"Target location(s): {location}\n"
        f"{skill_line}\n\n"
        f"RESUME (raw text — first 8000 chars):\n{resume_text[:8000]}"
    )


# ─────────────────────────────────────────────────
# RUN BUTTON
# ─────────────────────────────────────────────────
st.divider()
run_clicked = st.button("🚀 Run Agent", type="primary", disabled=not uploaded)

if run_clicked:
    # Reset state
    st.session_state.last_error = ""
    st.session_state.raw_final  = ""
    st.session_state.jobs_df    = None
    st.session_state.cover_doc  = None

    # Load and validate keys (sidebar input takes precedence)
    openai_key = openai_key_input.strip() if openai_key_input.strip() else os.environ.get("OPENAI_API_KEY", "").strip()
    tavily_key = tavily_key_input.strip() if tavily_key_input.strip() else os.environ.get("TAVILY_API_KEY", "").strip()

    if not openai_key:
        st.error("❌ OpenAI API key is missing. Please enter it in the sidebar or set it in the `.env` file.")
        st.stop()
    if not tavily_key:
        st.error("❌ Tavily API key is missing. Please enter it in the sidebar or set it in the `.env` file.")
        st.stop()

    # Inject tavily key into the mutable runtime dict — no module re-import needed
    _runtime["tavily_key"] = tavily_key

    try:
        resume_text = extract_text(uploaded)
        task = make_task_prompt(resume_text, skills_hint, target_title, target_location)

        agent = build_agent(openai_key, openai_model)

        initial_state = {
            "messages": [{"role": "user", "content": task}],
            "files": {"cover_letters.md": ""},
        }

        with st.spinner("🤖 Agent is searching for jobs and writing cover letters… (this may take 1–2 minutes)"):
            result = agent.invoke(initial_state)

        final_msgs = result.get("messages", [])
        final_text = (final_msgs[-1].content if final_msgs else "") or ""
        st.session_state.raw_final = final_text

        # Extract cover letters from virtual filesystem
        files    = result.get("files", {}) or {}
        cover_md = (files.get("cover_letters.md") or "").strip()
        st.session_state.cover_doc = md_to_docx(cover_md) if cover_md else None

        # Parse jobs
        raw_jobs  = extract_jobs_from_text(final_text)
        jobs_list = normalize_jobs(raw_jobs)
        st.session_state.jobs_df = pd.DataFrame(jobs_list) if jobs_list else None

        st.success("✅ Done! Scroll down to see your results.")

    except Exception as e:
        st.session_state.last_error = str(e)
        st.error(f"❌ Agent error: {e}")
        st.info("💡 Tip: Make sure both API keys are correct and your resume was uploaded.")

# ─────────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────────
if st.session_state.last_error:
    with st.expander("🔍 Debug info"):
        st.code(st.session_state.last_error, language="text")

st.divider()
col_jobs, col_dl = st.columns([3, 1])

with col_jobs:
    st.header("📋 Job Matches")
    if st.session_state.jobs_df is None or st.session_state.jobs_df.empty:
        st.info("No jobs to show yet. Upload your resume and click **Run Agent**.")
    else:
        df = st.session_state.jobs_df.copy()

        def as_link(u: str) -> str:
            u = u if isinstance(u, str) else ""
            return f'<a href="{u}" target="_blank">🔗 Apply</a>' if u else "—"

        if "Link" in df.columns:
            df["Link"] = df["Link"].apply(as_link)

        cols = [c for c in ["Company", "Title", "Location", "Link", "Good Match"] if c in df.columns]
        st.write(df[cols].to_html(escape=False, index=False), unsafe_allow_html=True)

with col_dl:
    st.header("📥 Download")
    if st.session_state.cover_doc:
        st.download_button(
            label="⬇️ Download cover_letters.docx",
            data=st.session_state.cover_doc,
            file_name="cover_letters.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_cover_letters",
        )
        st.caption("One personalised cover letter per job, ready to send.")
    else:
        st.info("Cover letters will appear here after the agent runs.")

# Raw agent output (collapsed by default)
if st.session_state.raw_final:
    with st.expander("🤖 Raw agent output"):
        st.text_area("Raw agent output", st.session_state.raw_final, height=300, label_visibility="collapsed")
